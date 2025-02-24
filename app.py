# import streamlit as st
# import pandas as pd
# import os
# from io import BytesIO
# from PyPDF2 import PdfReader
# from docx import Document
# import matplotlib.pyplot as plt
# import seaborn as sns





# # Set up the app
# st.set_page_config(page_title="Data Sweeper", layout="wide")

# # Custom CSS for Styling
# st.markdown("""
#     <style>
#         [data-testid="stSidebar"] {
#             background-color: #F5DEB3 !important;
#         }
#         [data-testid="stSidebarNav"] div {
#             background-color: #EED9A6;
#             padding: 12px;
#             border-radius: 8px;
#             margin-bottom: 5px;
#         }
#         [data-testid="stSidebarNav"] div:hover {
#             background-color: #E3C888;
#         }
#         .title {
#             color: #8B6914 !important;
#             font-weight: bold;
#             text-align: center;
#         }
#         div.stButton > button {
#             background-color: #EED9A6;
#             color: black;
#             border-radius: 8px;
#             font-size: 16px;
#             font-weight: bold;
#         }
#         div.stButton > button:hover {
#             background-color: #CDBA96;
#         }
#         .stAlert {
#             background-color: #F5F5DC !important;
#             color: black !important;
#             border-left: 6px solid #8B6914 !important;
#             border-radius: 8px;
#             padding: 12px;
#         }
#     </style>
# """, unsafe_allow_html=True)

# # Sidebar Navigation
# with st.sidebar:
#     st.markdown("## ⚙️ **Options**", unsafe_allow_html=True)
#     st.markdown("---")
#     page = st.radio("Navigation", ["🏠 Home", "📤 Upload & Convert"])

# # Main Content Area
# if page == "🏠 Home":
#     st.markdown("<h1 class='title'>📂 Data Sweeper - File Converter</h1>", unsafe_allow_html=True)
#     st.write("Transform your files between **CSV, Excel, PDF, and Word** formats with built-in **data cleaning and visualization!**")
#     st.markdown(
#         '<div class="stAlert">📢 <b>Go to "Upload & Convert" in the sidebar to process files!</b></div>',
#         unsafe_allow_html=True
#     )

# elif page == "📤 Upload & Convert":
#     st.markdown("<h1 class='title'>📤 Upload Your Files</h1>", unsafe_allow_html=True)
#     uploaded_files = st.file_uploader("Upload your files (CSV, Excel, or PDF):", 
#                                       type=["csv", "xlsx", "pdf"], 
#                                       accept_multiple_files=True)
#     success = False
    
#     if uploaded_files:
#         for file in uploaded_files:
#             file_ext = os.path.splitext(file.name)[-1].lower()
            
#             with st.spinner(f"🔄 Processing {file.name}..."):
#                 if file_ext == ".csv":
#                     df = pd.read_csv(file)
#                 elif file_ext == ".xlsx":
#                     df = pd.read_excel(file)
#                 elif file_ext == ".pdf":
#                     pdf_reader = PdfReader(file)
#                     text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
#                     doc = Document()
#                     doc.add_paragraph(text)
#                     buffer = BytesIO()
#                     doc.save(buffer)
#                     buffer.seek(0)
#                     file_name = file.name.replace(".pdf", ".docx")
#                     mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                     st.download_button(label=f"📥 Download {file.name} as Word", data=buffer, file_name=file_name, mime=mime_type)
#                     success = True
#                     continue

#             st.write(f"**📄 File Name:** {file.name}")
#             st.write(f"**📏 File Size:** {file.size / 1024:.2f} KB")
            
#             st.write("### 🏷️ Preview Data")
#             st.dataframe(df.head())
            
#             st.subheader("🧹 Data Cleaning Options")
#             if st.checkbox(f"Clean Data for {file.name}"):
#                 col1, col2 = st.columns(2)
                
#                 with col1:
#                     if st.button(f"🚮 Remove Duplicates from {file.name}"):
#                         df.drop_duplicates(inplace=True)
#                         st.write("✅ Duplicates Removed!")
                
#                 with col2:
#                     if st.button(f"🩹 Fill Missing Values for {file.name}"):
#                         numeric_cols = df.select_dtypes(include=['number']).columns
#                         df[numeric_cols] = df[numeric_cols].fillna(df[numeric_cols].mean())
#                         st.write("✅ Missing Values have been Filled")
            
#             st.subheader("📊 Select Columns to Convert")
#             columns = st.multiselect(f"Choose Columns for {file.name}", df.columns, default=df.columns)
#             df = df[columns]
            
#             st.subheader("📈 Data Visualization")
#             viz_option = st.selectbox(f"Choose visualization for {file.name}:", 
#                                       ["Bar Chart", "Line Chart", "Scatter Plot", "Pie Chart"])
            
#             if viz_option == "Bar Chart":
#                 st.bar_chart(df.select_dtypes(include='number'))
#             elif viz_option == "Line Chart":
#                 st.line_chart(df.select_dtypes(include='number'))
#             elif viz_option == "Scatter Plot":
#                 fig, ax = plt.subplots()
#                 sns.scatterplot(x=df.select_dtypes(include='number').columns[0], 
#                                 y=df.select_dtypes(include='number').columns[-1], 
#                                 data=df, ax=ax)
#                 st.pyplot(fig)
#             elif viz_option == "Pie Chart":
#                 fig, ax = plt.subplots()
#                 df[df.select_dtypes(include='number').columns[0]].value_counts().plot.pie(autopct='%1.1f%%', ax=ax)
#                 st.pyplot(fig)
            
#             st.subheader("🔄 Conversion Options")
#             conversion_type = st.radio(f"Convert {file.name} to:", ["CSV", "Excel"], key=file.name)
#             if st.button(f"🔄 Convert {file.name}"):
#                 buffer = BytesIO()
#                 if conversion_type == "CSV":
#                     df.to_csv(buffer, index=False)
#                     file_name = file.name.replace(file_ext, ".csv")
#                     mime_type = "text/csv"
#                 else:
#                     df.to_excel(buffer, index=False)
#                     file_name = file.name.replace(file_ext, ".xlsx")
#                     mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
#                 buffer.seek(0)
#                 st.download_button(label=f"📥 Download {file.name} as {conversion_type}", data=buffer, file_name=file_name, mime=mime_type)
#                 success = True
    
#     if success:
#         st.success("✅ All Files Processed Successfully! 🚀")
import streamlit as st
import pandas as pd
import os
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document
import matplotlib.pyplot as plt
import seaborn as sns

# Set up the app
st.set_page_config(page_title="Data Sweeper", layout="wide")

# Custom CSS for Styling
st.markdown("""
    <style>
        [data-testid="stSidebar"] { background-color: #F5DEB3 !important; }
        [data-testid="stSidebarNav"] div { background-color: #EED9A6; padding: 12px; border-radius: 8px; margin-bottom: 5px; }
        [data-testid="stSidebarNav"] div:hover { background-color: #E3C888; }
        .title { color: #8B6914 !important; font-weight: bold; text-align: center; }
        div.stButton > button { background-color: #EED9A6; color: black; border-radius: 8px; font-size: 16px; font-weight: bold; }
        div.stButton > button:hover { background-color: #CDBA96; }
        .stAlert { background-color: #F5F5DC !important; color: black !important; border-left: 6px solid #8B6914 !important; border-radius: 8px; padding: 12px; }
    </style>
""", unsafe_allow_html=True)

# Sidebar Navigation
with st.sidebar:
    st.markdown("## ⚙️ **Options**", unsafe_allow_html=True)
    st.markdown("---")
    page = st.radio("Navigation", ["🏠 Home", "📤 Upload & Convert"])

# Main Content Area
if page == "🏠 Home":
    st.markdown("<h1 class='title'>📂 Data Sweeper - File Converter</h1>", unsafe_allow_html=True)
    st.write("Transform your files between **CSV, Excel, PDF, and Word** formats with built-in **data cleaning and visualization!**")
    st.markdown('<div class="stAlert">📢 <b>Go to "Upload & Convert" in the sidebar to process files!</b></div>', unsafe_allow_html=True)

elif page == "📤 Upload & Convert":
    st.markdown("<h1 class='title'>📤 Upload Your Files</h1>", unsafe_allow_html=True)
    uploaded_files = st.file_uploader("Upload your files (CSV, Excel, or PDF):", type=["csv", "xlsx", "pdf"], accept_multiple_files=True)
    success = False
    
    if uploaded_files:
        for file in uploaded_files:
            file_ext = os.path.splitext(file.name)[-1].lower()
            
            with st.spinner(f"🔄 Processing {file.name}..."):
                if file_ext == ".csv":
                    df = pd.read_csv(file)
                elif file_ext == ".xlsx":
                    df = pd.read_excel(file)
                elif file_ext == ".pdf":
                    pdf_reader = PdfReader(file)
                    text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
                    
                    if not text.strip():
                        st.warning(f"⚠️ No readable text found in {file.name}.")
                        continue
                    
                    doc = Document()
                    doc.add_paragraph(text)
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    file_name = file.name.replace(".pdf", ".docx")
                    st.download_button(label=f"📥 Download {file.name} as Word", data=buffer, file_name=file_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    success = True
                    continue
                
                if df.empty:
                    st.warning(f"⚠️ {file.name} is empty! Please upload a valid file.")
                    continue
                
                st.write(f"**📄 File Name:** {file.name}")
                st.write(f"**📏 File Size:** {file.size / 1024:.2f} KB")
                st.write("### 🏷️ Preview Data")
                st.dataframe(df.head())
                
                st.subheader("🧹 Data Cleaning Options")
                if st.checkbox(f"Clean Data for {file.name}"):
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        if st.button(f"🚮 Remove Duplicates from {file.name}"):
                            df.drop_duplicates(inplace=True)
                            st.write("✅ Duplicates Removed!")
                    
                    with col2:
                        if st.button(f"🩹 Fill Missing Values for {file.name}"):
                            numeric_cols = df.select_dtypes(include=['number']).columns
                            df[numeric_cols] = df[numeric_cols].fillna(df[numeric_cols].mean())
                            st.write("✅ Missing Values have been Filled")
                
                st.subheader("📊 Select Columns to Convert")
                columns = st.multiselect(f"Choose Columns for {file.name}", df.columns, default=df.columns)
                df = df[columns]
                
                st.subheader("📈 Data Visualization")
                if not df.empty:
                    viz_option = st.selectbox(f"Choose visualization for {file.name}:", ["Bar Chart", "Line Chart", "Scatter Plot", "Pie Chart"])
                    
                    if viz_option == "Bar Chart":
                        st.bar_chart(df.select_dtypes(include='number'))
                    elif viz_option == "Line Chart":
                        st.line_chart(df.select_dtypes(include='number'))
                    elif viz_option == "Scatter Plot" and len(df.select_dtypes(include='number').columns) >= 2:
                        fig, ax = plt.subplots()
                        sns.scatterplot(x=df.select_dtypes(include='number').columns[0], 
                                        y=df.select_dtypes(include='number').columns[-1], 
                                        data=df, ax=ax)
                        st.pyplot(fig)
                    elif viz_option == "Pie Chart" and not df.select_dtypes(include='number').empty:
                        fig, ax = plt.subplots()
                        df[df.select_dtypes(include='number').columns[0]].value_counts().plot.pie(autopct='%1.1f%%', ax=ax)
                        st.pyplot(fig)
                
                st.subheader("🔄 Conversion Options")
                conversion_type = st.radio(f"Convert {file.name} to:", ["CSV", "Excel"], key=file.name)
                if st.button(f"🔄 Convert {file.name}"):
                    buffer = BytesIO()
                    if conversion_type == "CSV":
                        df.to_csv(buffer, index=False)
                        mime_type = "text/csv"
                    else:
                        df.to_excel(buffer, index=False)
                        mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    buffer.seek(0)
                    st.download_button(label=f"📥 Download {file.name} as {conversion_type}", data=buffer, file_name=f"converted_{file.name}", mime=mime_type)
                    success = True
    
    if success:
        st.success("✅ All Files Processed Successfully! 🚀")
